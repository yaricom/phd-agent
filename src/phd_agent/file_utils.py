"""
File utilities for writing essays in different formats.

This module provides functions to save essays in TXT, PDF, and DOCX formats.
"""

import logging
from pathlib import Path

from .models import Essay

logger = logging.getLogger(__name__)


def write_essay_txt(essay: Essay, output_path: str) -> bool:
    """Write an essay to a text file."""
    try:
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(f"Title: {essay.title}\n")
            f.write(f"Word Count: {essay.word_count}\n")
            f.write(f"Sources: {len(essay.sources)}\n")
            f.write(f"Created: {essay.created_at.strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write("=" * 50 + "\n\n")
            f.write(essay.content)
            f.write("\n\n" + "=" * 50 + "\n")
            f.write("SOURCES:\n")
            for i, source in enumerate(essay.sources, 1):
                f.write(f"{i}. {source.title} ({source.source_type.value})\n")
                if source.url:
                    f.write(f"   URL: {source.url}\n")
                f.write("\n")

        logger.info(f"Essay saved as TXT: {output_path}")
        return True

    except Exception as e:
        logger.error(f"Error writing TXT file: {e}")
        return False


def write_essay_pdf(essay: Essay, output_path: str) -> bool:
    """Write an essay to a PDF file."""
    # Try to import reportlab for PDF generation
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import (
            SimpleDocTemplate,
            Paragraph,
            Spacer,
            PageBreak,
        )
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
    except ImportError:
        logger.error("reportlab not installed. Install with: pip install reportlab")
        return False

    try:
        # Create PDF document
        doc = SimpleDocTemplate(output_path, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []

        # Title
        title_style = ParagraphStyle(
            "CustomTitle",
            parent=styles["Heading1"],
            fontSize=16,
            spaceAfter=30,
            alignment=1,  # Center
        )
        story.append(Paragraph(essay.title, title_style))
        story.append(Spacer(1, 20))

        # Metadata
        meta_style = ParagraphStyle(
            "Meta", parent=styles["Normal"], fontSize=10, textColor=colors.grey
        )
        story.append(Paragraph(f"Word Count: {essay.word_count}", meta_style))
        story.append(Paragraph(f"Sources: {len(essay.sources)}", meta_style))
        story.append(
            Paragraph(
                f"Created: {essay.created_at.strftime('%Y-%m-%d %H:%M:%S')}", meta_style
            )
        )
        story.append(Spacer(1, 20))

        # Content
        content_style = ParagraphStyle(
            "Content",
            parent=styles["Normal"],
            fontSize=11,
            spaceAfter=12,
            alignment=0,  # Left justify
        )

        # Split content into paragraphs
        paragraphs = essay.content.split("\n\n")
        for para in paragraphs:
            if para.strip():
                story.append(Paragraph(para.strip(), content_style))
                story.append(Spacer(1, 6))

        story.append(PageBreak())

        # Sources
        sources_style = ParagraphStyle(
            "Sources", parent=styles["Heading2"], fontSize=14, spaceAfter=20
        )
        story.append(Paragraph("SOURCES", sources_style))

        source_style = ParagraphStyle(
            "Source", parent=styles["Normal"], fontSize=10, spaceAfter=8
        )

        for i, source in enumerate(essay.sources, 1):
            source_text = f"{i}. {source.title} ({source.source_type.value})"
            if source.url:
                source_text += f"<br/>URL: {source.url}"
            story.append(Paragraph(source_text, source_style))
            story.append(Spacer(1, 4))

        # Build PDF
        doc.build(story)
        logger.info(f"Essay saved as PDF: {output_path}")
        return True

    except Exception as e:
        logger.error(f"Error writing PDF file: {e}")
        return False


def write_essay_docx(essay: Essay, output_path: str) -> bool:
    """Write an essay to a DOCX file."""
    # Try to import python-docx for DOCX generation
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        logger.error("python-docx not installed. Install with: pip install python-docx")
        return False

    try:
        # Create a document
        doc = Document()

        # Title
        title = doc.add_heading(essay.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        meta_para = doc.add_paragraph()
        meta_para.add_run(f"Word Count: {essay.word_count}\n")
        meta_para.add_run(f"Sources: {len(essay.sources)}\n")
        meta_para.add_run(f"Created: {essay.created_at.strftime('%Y-%m-%d %H:%M:%S')}")
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add some space
        doc.add_paragraph()

        # Content
        content_paragraphs = essay.content.split("\n\n")
        for para_text in content_paragraphs:
            if para_text.strip():
                para = doc.add_paragraph(para_text.strip())
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Add page break
        doc.add_page_break()

        # Sources
        doc.add_heading("SOURCES", level=1)

        for i, source in enumerate(essay.sources, 1):
            source_para = doc.add_paragraph()
            source_para.add_run(f"{i}. {source.title} ({source.source_type.value})")
            if source.url:
                source_para.add_run(f"\n   URL: {source.url}")

        # Save document
        doc.save(output_path)
        logger.info(f"Essay saved as DOCX: {output_path}")
        return True

    except Exception as e:
        logger.error(f"Error writing DOCX file: {e}")
        return False


def write_essay(essay: Essay, output_path: str, output_format: str = "auto") -> bool:
    """
    Write an essay to file in the specified format.

    Args:
        essay: The essay to write
        output_path: Output file path
        output_format: The output format to use ("txt", "pdf", "docx", or "auto" to detect from extension)

    Returns:
        bool: True if successful, False otherwise
    """
    if output_format == "auto":
        # Detect format from file extension
        ext = Path(output_path).suffix.lower()
        if ext == ".txt":
            output_format = "txt"
        elif ext == ".pdf":
            output_format = "pdf"
        elif ext == ".docx":
            output_format = "docx"
        else:
            # Default to txt if no extension or unknown extension
            output_format = "txt"
            if not output_path.endswith(".txt"):
                output_path += ".txt"

    # Ensure output directory exists
    output_dir = Path(output_path).parent
    output_dir.mkdir(parents=True, exist_ok=True)

    # Write based on format
    if output_format.lower() == "txt":
        return write_essay_txt(essay, output_path)
    elif output_format.lower() == "pdf":
        return write_essay_pdf(essay, output_path)
    elif output_format.lower() == "docx":
        return write_essay_docx(essay, output_path)
    else:
        logger.error(f"Unsupported format: {output_format}")
        return False


def get_supported_formats() -> list[str]:
    """Get a list of supported output formats."""
    formats = ["txt"]

    try:
        import reportlab  # noqa: F401

        formats.append("pdf")
    except ImportError:
        pass

    try:
        import docx  # noqa: F401

        formats.append("docx")
    except ImportError:
        pass

    return formats
